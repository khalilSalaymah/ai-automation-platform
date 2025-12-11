"""Knowledge base service for document processing."""

import uuid
import os
import tempfile
from typing import List, Dict, Any, Optional
from pathlib import Path
from loguru import logger

from .database import get_session
from .knowledge_models import (
    Document,
    DocumentChunk,
    AgentDocument,
    DocumentSource,
    DocumentStatus,
    DocumentSearchRequest,
    DocumentSearchResult,
)
from .embeddings import EmbeddingGenerator
from .vectorstore.pgvector import PGVectorStore
from .config import get_settings
from .errors import AgentFrameworkError

settings = get_settings()


class DocumentChunker:
    """Chunk documents into smaller pieces."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize chunker.

        Args:
            chunk_size: Maximum chunk size in characters
            chunk_overlap: Overlap between chunks in characters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_text(self, text: str) -> List[str]:
        """
        Chunk text into smaller pieces.

        Args:
            text: Text to chunk

        Returns:
            List of text chunks
        """
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind(".")
                last_newline = chunk.rfind("\n")
                break_point = max(last_period, last_newline)

                if break_point > start + self.chunk_size // 2:  # Only if reasonable
                    chunk = text[start : start + break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())
            start = end - self.chunk_overlap

        return chunks


class KnowledgeBaseService:
    """Service for managing knowledge base documents."""

    def __init__(self):
        """Initialize knowledge base service."""
        # EmbeddingGenerator uses Gemini embeddings (loads GEMINI_API_KEY from env)
        self.embedding_generator = EmbeddingGenerator()
        self.vector_store = PGVectorStore(
            uri=settings.database_url, table_name="document_embeddings"
        )
        self.chunker = DocumentChunker(chunk_size=1000, chunk_overlap=200)
        self.upload_dir = Path(tempfile.gettempdir()) / "knowledge_base_uploads"
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def _extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """
        Extract text from uploaded file.

        Args:
            file_path: Path to file
            file_type: MIME type or extension

        Returns:
            Extracted text
        """
        import mimetypes

        mime_type, _ = mimetypes.guess_type(file_path)
        if not mime_type:
            mime_type = file_type

        try:
            if mime_type in ["text/plain", "text/markdown"] or file_path.endswith(
                (".txt", ".md")
            ):
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()

            elif mime_type in [
                "application/pdf",
            ] or file_path.endswith(".pdf"):
                try:
                    import PyPDF2

                    text = ""
                    with open(file_path, "rb") as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    return text
                except ImportError:
                    logger.warning("PyPDF2 not installed, cannot extract PDF text")
                    raise AgentFrameworkError("PDF extraction requires PyPDF2")

            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ] or file_path.endswith(".docx"):
                try:
                    from docx import Document

                    doc = Document(file_path)
                    return "\n".join([para.text for para in doc.paragraphs])
                except ImportError:
                    logger.warning("python-docx not installed, cannot extract DOCX text")
                    raise AgentFrameworkError("DOCX extraction requires python-docx")

            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                raise AgentFrameworkError(f"Unsupported file type: {mime_type}")

        except Exception as e:
            logger.error(f"Error extracting text from file: {e}")
            raise AgentFrameworkError(f"Failed to extract text: {e}") from e

    def upload_document(
        self,
        name: str,
        file_content: bytes,
        file_type: str,
        user_id: Optional[str] = None,
        org_id: Optional[str] = None,
    ) -> str:
        """
        Upload and process a document.

        Args:
            name: Document name
            file_content: File content as bytes
            file_type: MIME type or extension
            user_id: User ID who uploaded
            org_id: Organization ID

        Returns:
            Document ID
        """
        doc_id = str(uuid.uuid4())
        file_path = self.upload_dir / f"{doc_id}_{name}"

        try:
            # Save file
            with open(file_path, "wb") as f:
                f.write(file_content)

            # Create document record
            with next(get_session()) as session:
                document = Document(
                    id=doc_id,
                    name=name,
                    source=DocumentSource.UPLOAD,
                    file_path=str(file_path),
                    file_type=file_type,
                    file_size=len(file_content),
                    status=DocumentStatus.PENDING,
                    user_id=user_id,
                    org_id=org_id,
                )
                session.add(document)
                session.commit()

            logger.info(f"Created document {doc_id}: {name}")
            return doc_id

        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise AgentFrameworkError(f"Failed to upload document: {e}") from e

    def process_document(self, document_id: str) -> None:
        """
        Process document: extract text, chunk, embed, and store.

        Args:
            document_id: Document ID
        """
        with next(get_session()) as session:
            document = session.get(Document, document_id)
            if not document:
                raise AgentFrameworkError(f"Document {document_id} not found")

            # Update status
            document.status = DocumentStatus.PROCESSING
            session.commit()

        try:
            # Extract text
            if document.file_path and os.path.exists(document.file_path):
                # For URL and Notion sources, text is stored in a temp file
                if document.source in [DocumentSource.URL, DocumentSource.NOTION]:
                    # Read text from temp file
                    with open(document.file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                else:
                    # For uploaded files, extract using file type
                    text = self._extract_text_from_file(document.file_path, document.file_type or "")
            else:
                raise AgentFrameworkError("Document file not found")

            # Chunk text
            chunks = self.chunker.chunk_text(text)
            total_chunks = len(chunks)

            # Update document with total chunks
            with next(get_session()) as session:
                document = session.get(Document, document_id)
                document.total_chunks = total_chunks
                session.commit()

            # Process chunks in batches
            batch_size = 100
            processed = 0

            for i in range(0, len(chunks), batch_size):
                batch_chunks = chunks[i : i + batch_size]

                # Generate embeddings
                embeddings = self.embedding_generator.generate(batch_chunks)

                # Create chunk records and prepare vector store data
                vector_ids = []
                vectors = []
                metadata_list = []

                for j, (chunk_text, embedding) in enumerate(zip(batch_chunks, embeddings)):
                    chunk_id = str(uuid.uuid4())
                    chunk_index = i + j

                    # Create chunk record
                    with next(get_session()) as session:
                        chunk = DocumentChunk(
                            id=chunk_id,
                            document_id=document_id,
                            chunk_index=chunk_index,
                            content=chunk_text,
                            vector_id=chunk_id,
                            metadata={"document_id": document_id, "chunk_index": chunk_index},
                        )
                        session.add(chunk)
                        session.commit()

                    vector_ids.append(chunk_id)
                    vectors.append(embedding)
                    metadata_list.append(
                        {
                            "document_id": document_id,
                            "chunk_id": chunk_id,
                            "chunk_index": chunk_index,
                        }
                    )

                # Store in vector store
                self.vector_store.add(vectors=vectors, metadata=metadata_list, ids=vector_ids)

                processed += len(batch_chunks)

                # Update progress
                with next(get_session()) as session:
                    document = session.get(Document, document_id)
                    document.processed_chunks = processed
                    session.commit()

            # Mark as completed
            with next(get_session()) as session:
                document = session.get(Document, document_id)
                document.status = DocumentStatus.COMPLETED
                session.commit()

            logger.info(f"Processed document {document_id}: {processed} chunks")

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            with next(get_session()) as session:
                document = session.get(Document, document_id)
                document.status = DocumentStatus.FAILED
                document.error_message = str(e)
                session.commit()
            raise

    def search_documents(
        self, request: DocumentSearchRequest, org_id: Optional[str] = None
    ) -> List[DocumentSearchResult]:
        """
        Search documents using semantic search.

        Args:
            request: Search request
            org_id: Organization ID for filtering

        Returns:
            List of search results
        """
        # Generate query embedding
        query_embedding = self.embedding_generator.generate([request.query])[0]

        # Build filter
        filter_dict = {}
        if request.document_ids:
            # Note: PGVector filter doesn't support IN queries directly
            # We'll filter after query
            pass

        # Query vector store
        results = self.vector_store.query(
            vector=query_embedding, top_k=request.top_k * 2, filter=filter_dict
        )

        # Get document info and filter
        document_ids_to_fetch = set()
        chunk_ids = []

        for result in results:
            chunk_id = result["id"]
            doc_id = result["metadata"].get("document_id")
            if doc_id:
                if request.document_ids and doc_id not in request.document_ids:
                    continue
                document_ids_to_fetch.add(doc_id)
                chunk_ids.append(chunk_id)

        # Fetch chunk details
        from sqlmodel import select

        with next(get_session()) as session:
            if chunk_ids:
                # SQLModel supports in_ on column attributes
                chunks = list(
                    session.exec(select(DocumentChunk).where(DocumentChunk.id.in_(chunk_ids)))
                )
            else:
                chunks = []
            
            if document_ids_to_fetch:
                documents = list(
                    session.exec(select(Document).where(Document.id.in_(document_ids_to_fetch)))
                )
            else:
                documents = []

            doc_map = {doc.id: doc for doc in documents}
            chunk_map = {chunk.id: chunk for chunk in chunks}

            # Build results
            search_results = []
            for result in results[: request.top_k]:
                chunk_id = result["id"]
                chunk = chunk_map.get(chunk_id)
                if not chunk:
                    continue

                doc = doc_map.get(chunk.document_id)
                if not doc:
                    continue

                # Filter by org_id if provided
                if org_id and doc.org_id != org_id:
                    continue

                # Filter by agent if provided
                if request.agent_name:
                    from sqlalchemy import and_
                    agent_docs = session.exec(
                        select(AgentDocument).where(
                            and_(
                                AgentDocument.agent_name == request.agent_name,
                                AgentDocument.document_id == doc.id,
                            )
                        )
                    ).first()
                    if not agent_docs:
                        continue

                search_results.append(
                    DocumentSearchResult(
                        chunk_id=chunk_id,
                        document_id=doc.id,
                        document_name=doc.name,
                        content=chunk.content,
                        score=result["score"],
                        chunk_index=chunk.chunk_index,
                        metadata=chunk.metadata,
                    )
                )

            return search_results

    def delete_document(self, document_id: str, org_id: Optional[str] = None) -> None:
        """
        Delete a document and all its chunks.

        Args:
            document_id: Document ID
            org_id: Organization ID for verification
        """
        with next(get_session()) as session:
            document = session.get(Document, document_id)
            if not document:
                raise AgentFrameworkError(f"Document {document_id} not found")

            if org_id and document.org_id != org_id:
                raise AgentFrameworkError("Document not found or access denied")

            # Get all chunks
            from sqlmodel import select

            chunks = list(
                session.exec(
                    select(DocumentChunk).where(DocumentChunk.document_id == document_id)
                ).all()
            )

            # Delete from vector store
            vector_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]
            if vector_ids:
                self.vector_store.delete(vector_ids)

            # Delete chunks
            for chunk in chunks:
                session.delete(chunk)

            # Delete agent associations
            agent_docs = list(
                session.exec(
                    select(AgentDocument).where(AgentDocument.document_id == document_id)
                ).all()
            )
            for agent_doc in agent_docs:
                session.delete(agent_doc)

            # Delete file if exists
            if document.file_path and os.path.exists(document.file_path):
                try:
                    os.remove(document.file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete file {document.file_path}: {e}")

            # Delete document
            session.delete(document)
            session.commit()

            logger.info(f"Deleted document {document_id}")

    def link_documents_to_agent(
        self, agent_name: str, document_ids: List[str], org_id: Optional[str] = None
    ) -> None:
        """
        Link documents to an agent.

        Args:
            agent_name: Agent name
            document_ids: List of document IDs
            org_id: Organization ID
        """
        from sqlmodel import select

        with next(get_session()) as session:
            # Verify documents exist and belong to org
            documents = list(
                session.exec(select(Document).where(Document.id.in_(document_ids)))
            )

            if org_id:
                documents = [doc for doc in documents if doc.org_id == org_id]

            if len(documents) != len(document_ids):
                raise AgentFrameworkError("Some documents not found or access denied")

            # Remove existing links for this agent
            existing = list(
                session.exec(
                    select(AgentDocument).where(AgentDocument.agent_name == agent_name)
                ).all()
            )
            for link in existing:
                session.delete(link)

            # Create new links
            for doc_id in document_ids:
                link = AgentDocument(
                    id=str(uuid.uuid4()),
                    agent_name=agent_name,
                    document_id=doc_id,
                    org_id=org_id,
                )
                session.add(link)

            session.commit()
            logger.info(f"Linked {len(document_ids)} documents to agent {agent_name}")

    def get_agent_documents(
        self, agent_name: str, org_id: Optional[str] = None
    ) -> List[Document]:
        """
        Get documents linked to an agent.

        Args:
            agent_name: Agent name
            org_id: Organization ID

        Returns:
            List of documents
        """
        from sqlmodel import select

        with next(get_session()) as session:
            query = (
                select(Document)
                .join(AgentDocument, Document.id == AgentDocument.document_id)
                .where(AgentDocument.agent_name == agent_name)
            )

            if org_id:
                query = query.where(AgentDocument.org_id == org_id)

            return list(session.exec(query).all())
